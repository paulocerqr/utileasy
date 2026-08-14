import sys
import tempfile
from pathlib import Path

import pymupdf
from django.test import SimpleTestCase, override_settings
from docx import Document
from PIL import Image as PillowImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Table, TableStyle

from apps.documents.models import DocumentConversion
from apps.documents.services import (
    DocumentConversionError,
    convert_document,
    inspect_docx,
    inspect_pdf,
    _run_bounded_process,
)


@override_settings(
    DOCUMENT_MAX_PAGES=200,
    DOCUMENT_MAX_UNCOMPRESSED_SIZE=20 * 1024 * 1024,
    DOCUMENT_CONVERSION_TIMEOUT_SECONDS=120,
    DOCUMENT_CONVERSION_MEMORY_LIMIT_MB=2048,
)
class DocumentConversionIntegrationTests(SimpleTestCase):
    def create_image(self, path):
        PillowImage.new("RGB", (80, 40), color=(35, 92, 140)).save(path, "PNG")

    def create_pdf_fixture(self, path, image_path):
        story = [
            Paragraph("Documento de referência Utileazy", getSampleStyleSheet()["Title"]),
            Image(str(image_path), width=80, height=40),
            Table([["Campo", "Valor"], ["Status", "Aprovado"]]),
        ]
        story[-1].setStyle(
            TableStyle(
                [
                    ("GRID", (0, 0), (-1, -1), 1, colors.black),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ]
            )
        )
        SimpleDocTemplate(str(path), pagesize=A4).build(story)

    def create_docx_fixture(self, path, image_path):
        document = Document()
        document.add_heading("Documento de referência Utileazy", level=1)
        document.add_picture(str(image_path))
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Campo"
        table.cell(0, 1).text = "Valor"
        table.cell(1, 0).text = "Status"
        table.cell(1, 1).text = "Aprovado"
        document.save(path)

    def test_docx_to_pdf_preserves_reference_text_and_table_content(self):
        with tempfile.TemporaryDirectory(prefix="utileazy-docx-pdf-") as temp_dir:
            root = Path(temp_dir)
            image_path = root / "reference.png"
            input_path = root / "reference.docx"
            self.create_image(image_path)
            self.create_docx_fixture(input_path, image_path)

            output_path, pages = convert_document(
                input_path,
                root,
                DocumentConversion.Operation.DOCX_TO_PDF,
            )

            self.assertEqual(pages, 1)
            self.assertTrue(output_path.read_bytes().startswith(b"%PDF-"))
            with pymupdf.open(output_path) as converted:
                text = "\n".join(page.get_text() for page in converted)
            self.assertIn("Documento de referência Utileazy", text)
            self.assertIn("Aprovado", text)

    def test_pdf_to_docx_preserves_reference_text_image_and_table(self):
        with tempfile.TemporaryDirectory(prefix="utileazy-pdf-docx-") as temp_dir:
            root = Path(temp_dir)
            image_path = root / "reference.png"
            input_path = root / "reference.pdf"
            self.create_image(image_path)
            self.create_pdf_fixture(input_path, image_path)

            output_path, pages = convert_document(
                input_path,
                root,
                DocumentConversion.Operation.PDF_TO_DOCX,
            )

            self.assertEqual(pages, 1)
            inspect_docx(output_path)
            converted = Document(output_path)
            text = "\n".join(paragraph.text for paragraph in converted.paragraphs)
            table_text = " ".join(
                cell.text
                for table in converted.tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertIn("Documento de referência Utileazy", text)
            self.assertIn("Aprovado", table_text)
            self.assertGreaterEqual(len(converted.inline_shapes), 1)

    def test_rejects_scanned_or_page_limit_pdf(self):
        with tempfile.TemporaryDirectory(prefix="utileazy-invalid-pdf-") as temp_dir:
            path = Path(temp_dir) / "image-only.pdf"
            document = pymupdf.open()
            document.new_page()
            document.save(path)
            document.close()

            with self.assertRaisesRegex(DocumentConversionError, "OCR"):
                inspect_pdf(path, require_text=True)
            with override_settings(DOCUMENT_MAX_PAGES=0):
                with self.assertRaisesRegex(DocumentConversionError, "limite"):
                    inspect_pdf(path, require_text=False)

    def test_rejects_password_protected_pdf(self):
        with tempfile.TemporaryDirectory(prefix="utileazy-protected-pdf-") as temp_dir:
            path = Path(temp_dir) / "protected.pdf"
            document = pymupdf.open()
            page = document.new_page()
            page.insert_text((72, 72), "Conteúdo protegido")
            document.save(
                path,
                encryption=pymupdf.PDF_ENCRYPT_AES_256,
                owner_pw="owner-password",
                user_pw="user-password",
            )
            document.close()

            with self.assertRaisesRegex(DocumentConversionError, "protegidos por senha"):
                inspect_pdf(path, require_text=True)

    def test_bounded_process_stops_after_timeout(self):
        with tempfile.TemporaryDirectory(prefix="utileazy-timeout-") as temp_dir:
            with override_settings(DOCUMENT_CONVERSION_TIMEOUT_SECONDS=1):
                with self.assertRaisesRegex(DocumentConversionError, "tempo limite"):
                    _run_bounded_process(
                        [sys.executable, "-c", "import time; time.sleep(5)"],
                        temp_dir,
                        "Falha inesperada.",
                    )
